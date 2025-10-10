import contextlib
import json
import os
import re
from copy import deepcopy
from datetime import date
from typing import Any, Optional

from docx import Document
from docx.document import Document as DocumentClass
from docx.opc.oxml import BaseOxmlElement
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.parts.story import StoryPart
from docx.text.paragraph import Paragraph
from docxtpl import DocxTemplate


def save_raw_data(data: dict[str, list[dict[str, str | list[str]]]], save_location: str, filename: str) -> None:
    """Save compiled news data as JSON

    Args:
        data (dict[str, list[dict[str, str | list[str]]]]): the compiled news data
        save_location (str): the folder location where file will be saved
        filename (str): as name suggests
    """
    if not os.path.exists(save_location):
        os.makedirs(save_location)
    with open(os.path.join(save_location, filename), "w") as f:
        json.dump(data, f)


def save_processsed_data(data: dict[str, list[dict[str, str | list[str]]]], save_location: str, filename: str, template: str) -> None:
    """Save compiled news data as DOCX

    Args:
        data (dict[str, list[dict[str, str | list[str]]]]): the compiled news data
        save_location (str): the folder location where file will be saved
        filename (str): as name suggests
        template (str): .DOCX template filepath to use for compiling news data
    """
    if not os.path.exists(save_location):
        os.makedirs(save_location)

    filepath = os.path.join(save_location, filename)
    tmp_path = filepath.replace(".docx", ".tmp.docx")

    tpl = DocxTemplate(template_file=template)
    tpl.render(context={"news": data, "header_date": date.today().strftime("%A, %d %B, %Y")})
    tpl.save(tmp_path)

    doc = Document(tmp_path)
    _walk_and_replace(doc)
    doc.save(filepath)

    # cleanup tmp
    with contextlib.suppress(OSError):
        os.remove(tmp_path)


def _add_hyperlink(document_part: StoryPart, url: str, text: str, rPr_copy: Optional[Any] = None) -> BaseOxmlElement:
    """
    Build a <w:hyperlink> element with a child <w:r> containing rPr_copy (if given)
    and a <w:t> text node for `text`.
    Returns the hyperlink OxmlElement.
    """
    # create relationship id for hyperlink target
    r_id = document_part.relate_to(
        url,
        reltype="http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    run_elem = OxmlElement("w:r")
    # attach rPr copy if available
    if rPr_copy is not None:
        # deepcopy to avoid reusing the same node
        run_elem.append(deepcopy(rPr_copy))

    # add text node
    text_node = OxmlElement("w:t")
    if text.startswith(" ") or text.endswith(" "):
        text_node.set(qn("xml:space"), "preserve")
    text_node.text = text
    run_elem.append(text_node)

    hyperlink.append(run_elem)
    return hyperlink


def _process_paragraph_replace_urls(paragraph: Paragraph) -> None:
    # sourcery skip: low-code-quality
    """
    Replace URL substrings inside a paragraph with actual hyperlinks.
    Preserves surrounding text segments.
    """
    # crude-but-useful URL regex
    URL_RE = re.compile(r"(https?://[^\s\)\]\>]+)")

    # gather runs and their text
    runs = list(paragraph.runs)
    if not runs:
        return

    run_texts = [r.text or "" for r in runs]
    full_text = "".join(run_texts)
    if not URL_RE.search(full_text):
        return  # nothing to do

    # Build a per-character mapping to original run index
    char_run_idx: list[int] = []
    for idx, txt in enumerate(run_texts):
        char_run_idx.extend([idx] * len(txt))

    # Find URL matches in the concatenated text
    matches = list(URL_RE.finditer(full_text))
    # Build segments: (text_segment, is_url, src_run_idx)
    parts: list[tuple[str, bool, int | None]] = []
    last_end = 0
    for m in matches:
        if m.start() > last_end:
            # non-url segment
            seg_text = full_text[last_end : m.start()]
            src_idx = char_run_idx[last_end] if last_end < len(char_run_idx) else None
            parts.append((seg_text, False, src_idx))
        # url segment
        url_text = m.group(0)
        src_idx = char_run_idx[m.start()] if m.start() < len(char_run_idx) else None
        parts.append((url_text, True, src_idx))
        last_end = m.end()
    if last_end < len(full_text):
        seg_text = full_text[last_end:]
        src_idx = char_run_idx[last_end] if last_end < len(char_run_idx) else None
        parts.append((seg_text, False, src_idx))

    # Remove all original runs
    for r in list(paragraph.runs):
        paragraph._p.remove(r._element)

    # Recreate runs & hyperlinks, copying rPr from the original source run when available
    for seg_text, is_url, src_idx in parts:
        if seg_text == "":
            continue
        if is_url:
            # create hyperlink element with rPr copied from the source run if possible
            rPr_copy = None
            if src_idx is not None and 0 <= src_idx < len(runs):
                src_r = runs[src_idx]
                # src_r._element.rPr may be missing
                rPr_copy = getattr(src_r._element, "rPr", None)
            hyperlink_el = _add_hyperlink(paragraph.part, seg_text, seg_text, rPr_copy=rPr_copy)
            paragraph._p.append(hyperlink_el)
        else:
            # non-url: create a run and copy formatting from source run if available
            # Add run via paragraph.add_run then attach rPr copy if exists
            new_run = paragraph.add_run(seg_text)
            if src_idx is not None and 0 <= src_idx < len(runs):
                src_r = runs[src_idx]
                src_rPr = getattr(src_r._element, "rPr", None)
                if src_rPr is not None:
                    # insert a deep-copied rPr as the first child of new run element
                    new_run_elm = new_run._element
                    # remove existing default rPr if any, then insert the copied one
                    existing_rPr = new_run_elm.find(qn("w:rPr"), namespaces=new_run_elm.nsmap)
                    if existing_rPr is not None:
                        new_run_elm.remove(existing_rPr)
                    new_run_elm.insert(0, deepcopy(src_rPr))


def _walk_and_replace(doc: DocumentClass) -> None:
    # Paragraphs in body
    for para in doc.paragraphs:
        _process_paragraph_replace_urls(para)
